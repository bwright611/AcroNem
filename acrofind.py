import re
import sys
from docx import Document
from docx.shared import Inches
import openpyxl
from openpyxl import load_workbook
import json
import os
import csv
import argparse

#!/usr/bin/env python3
"""
acrofind.py - simple acronym finder
Usage:
    python acrofind.py [paths...] [-r] [--min LEN] [--json] [-v]
If no paths are given, reads from stdin.
"""
'''
ACRO_RE = re.compile(
        r"""(
                (?:[A-Z]\.){2,}         # dotted acronyms like U.S.A.
                |
                \b(?:(?:[A-Z]{2,}|(?:[a-z]*[A-Z][a-z]*){2,}|[A-Z](?:[-&][A-Z]+)+)\d*|\d+[A-Z]|[A-Z]\d+)\b  # consecutive caps like NASA or X-99
        )""",
        re.VERBOSE,
)
'''
ACRO_RE = re.compile(r"""
    (                                   # ── whole match ──
        (?:[A-Z]\.){2,}                 #   dotted acronyms:  U.S.A.,  N.A.S.A.
      | \b
        (?:                             #   ── “real” acronyms ──
            \d+[A-Z]{1,}                #   2FA , 3GPP               (digits → caps)
          | [A-Z]\d+                    #   C5I  (caps → digits)
          | [A-Z]\d+[A-Z]                 # C5I, X9Y   (cap‑digit‑cap)
          | [A-Z]{2,}                   #   ARROW, ATS, US, WAN …   (plain all‑caps)
          | [A-Z]+(?:[-&/][A-Z]+)+      #   USB‑C, TSM‑X, O&M, AN/PRC (caps separated by – & /)
          # | [a-z]+/[a-z]+               #   ft/s                    (lower‑case slash)
          # | [A-Z]{2,}[a-z]+             #   MMWave, MIL‑SPEC        (caps‑then‑lower)
          # | [A-Za-z]*[A-Z][A-Za-z]*[a-z][A-Za-z]*   # mixed‑case with at least one lower‑case letter
          # ---- optional explicit lower‑case units ----
          | \b(?:bps|ft|km|kpi|mb|gbps|ft/s)\b   # uncomment & edit if you need them
        )
        \b
    )
""", re.VERBOSE)


def normalize_acronym(s: str) -> str:
        """Normalize an acronym string by stripping trailing periods only.
        Plural-handling is performed after collecting all matches so we
        only remove a trailing 's' when the singular form appears elsewhere.
        """
        return s.rstrip('.')


def is_likely_roman(s: str) -> bool:
        """Return True if `s` looks like a simple Roman numeral we want to
        exclude (common enumerations like I, II, III, IV, V, VI, VII, VIII,
        IX, X, XI, XII). This intentionally only matches numerals using
        the letters I, V, X up to length 4 to avoid filtering valid
        acronyms like `MD` or `NV`.
        """
        if not s:
                return False
        t = s.rstrip('.').upper()
        return bool(re.fullmatch(r'[IVX]{1,4}', t))



def createWordDocument(acronym_list: list[str], output_path: str, definitions: dict[str, str] | None = None) -> None:
        """Create a Word document with the list of acronyms in a two-column table (Acronym | Definition).
        If `definitions` is provided it will be used to fill the Definition column for matching acronyms.
        """
        doc = Document()
        doc.add_heading('Acronym List', 0)

        # Create a table with a header row and two columns: Acronym and Definition
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Acronym'
        hdr_cells[1].text = 'Definition'

        # Make header text bold
        for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                                run.bold = True

        # Add one row per acronym, using provided definitions when available
        for acronym in acronym_list:
                row_cells = table.add_row().cells
                row_cells[0].text = acronym
                row_cells[1].text = (definitions.get(acronym, '') if definitions else '')

        # Best-effort column widths (may be ignored by some Word viewers)
        try:
                table.columns[0].width = Inches(1.5)
                table.columns[1].width = Inches(4)
        except Exception:
                pass

        doc.save(output_path)
"""
 def find_acronyms_in_text(text: str, min_len: int = 2) -> Counter:
        #Return counter of acronyms found in text.
        matches = (m.group(0) for m in ACRO_RE.finditer(text))
        normalized = (m.rstrip(".") for m in matches) 
        filtered = (m for m in normalized if len(m.replace(".", "")) >= min_len)
        return Counter(filtered)
"""

def importWordDocx(file_path: str) -> str:
        """Extract text from a .docx file."""
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
                full_text.append(para.text)
        return "\n".join(full_text)


def importAcronymList(file_path: str) -> dict[str, str]:
    """Load a master acronym list from .xlsx/.csv/.json/.docx into a dict acro->definition."""
    result: dict[str, str] = {}
    if not file_path or not os.path.exists(file_path):
            return result

    ext = os.path.splitext(file_path)[1].lower()
    if ext in ('.xlsx', '.xls'):
            try:
                    wb = load_workbook(file_path, read_only=True, data_only=True)
                    ws = wb.active
                    for row in ws.iter_rows(min_row=1, values_only=True):
                            if not row or not row[0]:
                                    continue
                            acro_raw = str(row[0]).strip()
                            definition_raw = str(row[1]).strip() if len(row) > 1 and row[1] is not None else ''
                            if not acro_raw or not definition_raw:
                                    # skip incomplete rows
                                    continue
                            acro = normalize_acronym(acro_raw)
                            result[acro] = definition_raw
            except Exception as e:
                    print(f"Error reading Excel {file_path}: {e}", file=sys.stderr)
    elif ext == '.csv':
            with open(file_path, newline='', encoding='utf-8') as f:
                    reader = csv.reader(f)
                    for row in reader:
                            if not row or not row[0].strip():
                                    continue
                            acro_raw = row[0].strip()
                            definition_raw = row[1].strip() if len(row) > 1 else ''
                            if not acro_raw or not definition_raw:
                                    continue
                            acro = normalize_acronym(acro_raw)
                            result[acro] = definition_raw
    elif ext == '.json':
            with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    if isinstance(data, dict):
                            for k, v in data.items():
                                    if not k:
                                            continue
                                    v_str = '' if v is None else str(v).strip()
                                    if not v_str:
                                            continue
                                    result[normalize_acronym(k)] = v_str
                    elif isinstance(data, list):
                            for item in data:
                                    if isinstance(item, dict):
                                            k = item.get('acro') or item.get('acronym') or item.get('key')
                                            v = item.get('def') or item.get('definition') or ''
                                            if not k or not v:
                                                    continue
                                            v_str = str(v).strip()
                                            if not v_str:
                                                    continue
                                            result[normalize_acronym(k)] = v_str
    elif ext == '.docx':
            try:
                    doc = Document(file_path)
                    for table in doc.tables:
                            for row in table.rows:
                                    cells = [c.text.strip() for c in row.cells]
                                    if not cells or not cells[0].strip():
                                            continue
                                    if len(cells) < 2 or not cells[1].strip():
                                            continue
                                    acro = normalize_acronym(cells[0])
                                    definition = cells[1].strip()
                                    result[acro] = definition
            except Exception as e:
                    print(f"Error reading DOCX {file_path}: {e}", file=sys.stderr)

    return result

def main(argv: list[str] | None = None) -> int:
        parser = argparse.ArgumentParser(description='Find acronyms and build an acronym list docx')
        parser.add_argument('--master', '-m', help='path to master acronym file (.xlsx/.csv/.json/.docx)')
        parser.add_argument('--out', '-o', help='output Word filename', default='Acronym_List.docx')
        args = parser.parse_args(argv) if argv is not None else parser.parse_args()

        document = Document('TARCES Technical Volume - Test.docx')
        # document.save('newTarces.docx')
        
        acronym_list = []
        found_in_paragraphs = []

        print("\nSearching in paragraphs:")
        for paragraph in document.paragraphs:
                for m in ACRO_RE.findall(paragraph.text):
                        norm = normalize_acronym(m)
                        if is_likely_roman(norm):
                                continue
                        acronym_list.append(norm)

        print("\nSearching in tables...")
        for table in document.tables:
                for row in table.rows:
                                for cell in row.cells:
                                        for m in ACRO_RE.findall(cell.text):
                                                norm = normalize_acronym(m)
                                                if is_likely_roman(norm):
                                                        continue
                                                acronym_list.append(norm)
        exclude_list = ['1366E', '141B', '175F', '175X', '181D', '220D', '31000B', '400S', '881F', 'A001', 'A002', 'A004', 'A006', 'A007', 'A010', 'A011', 'A012', 'A013', 'A022', 'A026', 'A027', 'COMMAND', 'COMMUNICATIONS', 'COMPUTERS', 'CONOPS', 'CONTROL', 'INTELLIGENCE', 'TACTICAL', 'TECHNICAL', 'SYSTEMS', 'REMOTE']
        # Deduplicate, but if both singular and a plural-with-lowercase-s
        # appear, prefer the singular and drop the plural (e.g., CDRLs -> CDRL).
        all_set = set(acronym_list)
        to_remove = set()
        for a in list(all_set):
                if len(a) > 1 and a.endswith('s'):
                        base = a[:-1]
                        if base in all_set:
                                to_remove.add(a)
        all_set.difference_update(to_remove)

        deduped_acronym_list = list(all_set)
        sorted_acronym_list = sorted(deduped_acronym_list, key=lambda v: v.upper())
        sorted_acronym_list = [x for x in sorted_acronym_list if x not in exclude_list]
        print(sorted_acronym_list)
        print(f"The number of acronyms is: {len(sorted_acronym_list)}")

        # createWordDocument(sorted_acronym_list, 'Acronym_List.docx')

        # Use CLI-provided master path if given; for testing overwrite it
        master_path = args.master or 'ISS_Master_List.xlsx'
        # Overwrite for testing (explicit request): always use the function's path
        master_path = 'ISS_Master_List.xlsx'
        if args.master:
                print(f"Note: overriding provided --master {args.master} for testing; using {master_path}")

        if not os.path.exists(master_path):
                print(f"Master file not found: {master_path}", file=sys.stderr)
                master = {}
        else:
                master = importAcronymList(master_path)

        out_path = 'master_acronyms.txt'
        try:
                with open(out_path, 'w', encoding='utf-8') as f:
                        f.write(json.dumps(master, indent=2, ensure_ascii=False))
                print(f"Wrote master dict to {out_path}")
        except Exception as e:
                print(f"Error writing master dict to {out_path}: {e}", file=sys.stderr)

        # Build mapping of found acronyms to their definitions (empty if not found)
        acro_defs = {a: master.get(a, '') for a in sorted_acronym_list}

        # Create the Word document with definitions populated when available
        try:
                createWordDocument(sorted_acronym_list, args.out, acro_defs)
                print(f"Wrote {args.out} with definitions (where available)")
        except Exception as e:
                print(f"Error creating {args.out}: {e}", file=sys.stderr)

# testing comments



                

if __name__ == "__main__":
        sys.exit(main())
